import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"

TITLE_FONT = "黑体"
BODY_FONT = "宋体"
WESTERN_FONT = "Times New Roman"
TITLE_SIZE = Pt(16)  # 三号
BODY_SIZE = Pt(12)  # 小四
LINE_SPACING = Pt(22)
FIRST_LINE_INDENT = Pt(24)


def set_run_font(run, east_asia=BODY_FONT, size=BODY_SIZE, bold=False):
    run.font.name = WESTERN_FONT
    run.font.size = size
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), WESTERN_FONT)
    r_fonts.set(qn("w:hAnsi"), WESTERN_FONT)
    r_fonts.set(qn("w:cs"), WESTERN_FONT)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_format(paragraph, first_line=True, align=None):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = LINE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = FIRST_LINE_INDENT
    else:
        fmt.first_line_indent = Pt(0)
    if align is not None:
        paragraph.alignment = align


def add_title(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=TITLE_FONT, size=TITLE_SIZE, bold=True)


def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line=False)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=TITLE_FONT, size=TITLE_SIZE, bold=True)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line=True)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_no_indent_body(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line=False)
    run = paragraph.add_run(text)
    set_run_font(run)


def style_cell(cell, header=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        set_paragraph_format(paragraph, first_line=False)
        for run in paragraph.runs:
            set_run_font(run, east_asia=TITLE_FONT if header else BODY_FONT, size=BODY_SIZE, bold=header)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, value in enumerate(headers):
        hdr[idx].text = value
        style_cell(hdr[idx], header=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            style_cell(cells[idx], header=False)
    doc.add_paragraph()


def set_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = WESTERN_FONT
    normal.font.size = BODY_SIZE
    normal._element.rPr.rFonts.set(qn("w:ascii"), WESTERN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), WESTERN_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), WESTERN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)


def create_project_manual():
    doc = Document()
    set_document_styles(doc)

    add_title(doc, "校园 AI 互助撮合平台项目说明书")
    add_no_indent_body(doc, "提交材料：项目说明书")
    add_no_indent_body(doc, "版本日期：2026 年 5 月 30 日")

    add_heading(doc, "一、项目概述")
    add_body(doc, "校园 AI 互助撮合平台面向高校学生在竞赛组队、课程项目、大创申报、技术求助和经验交流中的协作需求，提供从需求表达、AI 辅助整理、智能匹配、主动申请到站内沟通的闭环服务。平台不是传统信息发布论坛，而是以 AI Agent 为核心的校园协作工作台，帮助用户把模糊想法转化为结构化需求，并进一步找到合适的协作者。")
    add_body(doc, "项目的核心价值在于降低校园协作的信息不对称。发布方可以更清楚地描述自己需要什么，参与方可以更快发现适合自己的机会，双方在匹配或申请后能够直接进入沟通，减少微信群、朋友圈和线下询问带来的低效与遗漏。")

    add_heading(doc, "二、建设背景与问题")
    add_body(doc, "在校园场景中，学生经常会遇到“想参赛但找不到队友”“有项目但不知道谁合适”“想加入机会但不知道哪里有需求”等问题。传统方式依赖熟人推荐或社群发帖，信息分散、描述不完整、状态不可追踪，容易出现需求无人响应、重复沟通、候选人不匹配等情况。")
    add_body(doc, "本项目以黑客松、算法竞赛、蓝桥杯、大创、课程项目和技术求助等高频场景为切入点，通过用户画像、技能标签、需求语义、行为记录和 Agent 对话上下文，构建一个可演示、可运行、可扩展的校园 AI 协作入口。")

    add_heading(doc, "三、目标用户")
    add_body(doc, "第一类用户是需求发布方，包括正在寻找队友、技术支持、项目展示协作者或导师建议的学生。第二类用户是参与方，包括希望发现比赛、项目、求助机会并主动申请加入的学生。第三类用户是平台管理或演示人员，他们需要稳定查看需求、匹配、申请和消息状态，确保协作流程可追踪。")

    add_heading(doc, "四、核心功能")
    add_table(
        doc,
        ["功能模块", "主要能力"],
        [
            ["用户与画像", "支持登录、资料编辑、技能标签、院校专业等扩展信息，用于匹配和 AI 个性化生成。"],
            ["需求发布", "支持手动发布组队、求助、技能交换等需求，并可选择单人或多人模式。"],
            ["AI 辅助创作", "支持 AI 生成和润色需求描述，生成时读取当前用户画像，并遵守单人/多人选择约束。"],
            ["Agent 工作台", "支持自然语言追问、文件分析、生成需求草稿、确认发布、反向推荐已有需求和起草申请。"],
            ["智能匹配", "基于技能标签、需求语义和用户画像生成候选人列表，并给出匹配理由和候选人对比。"],
            ["主动申请", "参与方可在需求详情页提交申请，发布方可接受或拒绝，申请方可查看状态。"],
            ["消息沟通", "匹配或申请后可进入站内消息页，围绕具体需求继续沟通。"],
            ["草稿保护", "发布表单、申请草稿和私信草稿支持本地持久化，切换页面后不会丢失。"],
        ],
    )

    add_heading(doc, "五、典型业务流程")
    add_body(doc, "发布方流程为：用户登录平台后进入 Agent 工作台，通过聊天或上传比赛通知描述需求；Agent 根据需求类型、标题、描述、技能方向和人数模式进行追问；信息充分后生成可发布草稿；用户确认发布；系统启动匹配；发布方查看候选人、比较匹配理由、选择候选人并进入消息沟通。")
    add_body(doc, "参与方流程为：用户可以浏览需求广场，也可以直接告诉 Agent 自己想参加的方向，例如算法比赛或黑客松；Agent 识别为发现已有需求后推荐相关机会；用户进入需求详情页，让 AI 起草申请文案并提交；发布方处理申请后，申请方在我的申请页面看到已通过、未通过或待处理状态。")
    add_body(doc, "单人需求和多人需求的状态逻辑不同。单人需求完成配对后，未被选择的待处理申请会进入终态，避免出现发布方已经选定人选但其他申请方仍看到待处理的问题。多人需求则允许保留多个被选中用户，适合团队协作类项目。")

    add_heading(doc, "六、AI 能力说明")
    add_body(doc, "平台将大模型能力拆分为 Prompt、Skill、Agent 和事件链路。PromptRegistry 管理语义路由、文件分析、需求生成、消息起草、匹配解释等提示词模板；SkillRegistry 封装语义路由、标签提取、Embedding、向量匹配、文件读取、内容审核和上下文摘要等能力；多个 Agent 共同完成意图识别、任务规划、需求创建、语义搜索、重排解释和匹配提醒。")
    add_body(doc, "用户画像是 AI 个性化生成的重要基础。平台在生成需求描述、润色内容、起草申请和私信时，会结合用户的技能标签、简介、学校学院专业等扩展信息、偏好画像、历史需求风格和近期行为统计，使生成内容更贴近当前用户，而不是输出通用模板。")
    add_body(doc, "语义路由能力用于区分“发布需求”和“寻找已有需求”。例如用户说“我想打算法比赛，帮我看看有没有这方面的组队需求”时，系统应识别为发现已有机会，而不是直接进入发布流程；当用户说“帮我发布一个蓝桥杯组队需求”时，系统会继续追问技能方向和人数模式，避免过早发布粗糙需求。")

    add_heading(doc, "七、系统特色")
    add_body(doc, "第一，平台支持双边撮合，不只服务发布方，也服务想找机会的参与方。第二，Agent 不是单纯聊天，而是能把自然语言、文件内容和用户画像转化为可执行业务流程。第三，匹配结果具有可解释性，用户不仅能看到候选人，还能看到推荐理由和候选人对比。第四，平台重视演示稳定性，对权限、重复提交、删除关联数据、申请状态流转和草稿丢失等问题进行了专项修复。")

    add_heading(doc, "八、演示场景")
    add_body(doc, "推荐演示主线是：使用 alice 登录，进入智能助手，输入或上传黑客松需求，Agent 追问并生成草稿，用户确认发布，进入匹配结果页查看候选人，选择 iris 或起草私信，进入消息页沟通。随后切换到 iris 账号，在我的需求页面查看自己被选中的需求并继续互动。")
    add_body(doc, "推荐演示副线是：使用 bob 登录，让 Agent 查找算法比赛相关已有需求，进入需求详情页，AI 起草申请内容并提交；发布方处理后，bob 在我的申请页面查看状态变化。该副线可体现平台对“找机会”场景的支持。")

    add_heading(doc, "九、测试与稳定性")
    add_body(doc, "赛前已围绕登录、发布需求、Agent 草稿、文件上传、匹配、申请、消息、删除需求、系统设置 API Key 和页面切换草稿保存等主链路进行测试。最近一次验证中，后端契约测试 tests.test_project_contracts 共 48 项通过，Agent 烟测 tests.test_agent_smoke 共 19 项通过，后端 compileall 通过，前端 npm run build 通过，后端健康检查 /api/health 返回 ok。")
    add_body(doc, "当前已知非阻塞项包括前端构建中的 Rolldown 注释警告和 vendor chunk size warning，以及外部模型服务不可达时 AI 生成质量可能降级。即使模型临时不可用，平台仍保留需求、匹配、申请和消息等核心业务链路。")

    add_heading(doc, "十、后续规划")
    add_body(doc, "后续可继续推进更完整的自动执行器、长期记忆与个人偏好学习、WebSocket 实时协作、小程序端、合作反馈闭环和匹配质量评估。草稿持久化目前采用浏览器 localStorage，适合演示和本地使用，未来可升级为后端草稿表，实现跨设备连续体验。")

    out = OUT_DIR / "项目说明书_校园AI互助撮合平台_提交版.docx"
    doc.save(out)
    return out


def create_architecture_doc():
    doc = Document()
    set_document_styles(doc)

    add_title(doc, "校园 AI 互助撮合平台技术架构文档")
    add_no_indent_body(doc, "提交材料：技术架构文档")
    add_no_indent_body(doc, "版本日期：2026 年 5 月 30 日")

    add_heading(doc, "一、架构总览")
    add_body(doc, "平台采用前后端分离架构。前端基于 Vue 3、TypeScript、Vite、Pinia、Vue Router 和 Element Plus 构建，负责用户界面、路由、状态管理和 API 调用。后端基于 FastAPI、SQLAlchemy Async、SQLite 和 aiosqlite 构建，负责认证、需求、匹配、申请、消息、Agent、设置和文件处理等业务服务。AI 能力通过模型客户端、PromptRegistry、SkillRegistry、Agent 编排器和本地模型适配器组合实现。")
    add_body(doc, "整体架构可概括为四层：用户界面层负责交互表达；API 接入层负责身份认证和 REST 接口；业务与 Agent 层负责需求发布、匹配、申请、消息和 AI 推理；数据与模型层负责 SQLite 数据持久化、文件存储、Embedding、Rerank 和外部大模型调用。")

    add_table(
        doc,
        ["层级", "技术与职责"],
        [
            ["用户界面层", "Vue 3 + Element Plus，提供登录、需求广场、发布需求、我的需求、匹配结果、消息、Agent 工作台和系统设置页面。"],
            ["API 接入层", "FastAPI Router，提供 /api/auth、/api/needs、/api/agent、/api/messages、/api/profile、/api/settings 等接口。"],
            ["业务服务层", "NeedService、MatchEngine、NeedApplicationService、MessageService、ProfileService 等服务实现核心业务规则。"],
            ["Agent 能力层", "IntentAnalyzerAgent、PlannerAgent、NeedCreatorAgent、SemanticSearchAgent、RerankAgent、MatchWatcherAgent 等协同工作。"],
            ["数据与模型层", "SQLite 保存用户、需求、申请、匹配、消息和 Agent 会话；DeepSeek 负责生成类任务；Qwen3 Embedding/Reranker 负责本地语义召回与精排。"],
        ],
    )

    add_heading(doc, "二、前端架构")
    add_body(doc, "前端项目位于 frontend 目录，使用 Vue 3 的组合式 API 编写页面逻辑，使用 Pinia 管理登录态、需求数据和 Agent 状态，使用 Vue Router 管理页面跳转。Vite 提供开发服务和构建能力，Axios 封装 HTTP 请求，Element Plus 提供基础组件。")
    add_body(doc, "主要页面包括 AgentView、NeedPlazaView、NeedCreateView、NeedDetailView、NeedManageView、MatchResultView、MyApplicationsView、MessagesView、ProfileView 和 SettingsView。各页面围绕用户的发布、发现、申请、匹配和沟通流程组织，减少跨页面状态断裂。")
    add_body(doc, "前端最近补充了 persistentDrafts 工具，将需求发布表单、申请文案和匹配私信草稿保存到 localStorage。草稿键按当前用户和需求 ID 区分，避免账号切换或不同需求之间串数据。成功发布或成功申请后，对应草稿会自动清除。")

    add_heading(doc, "三、后端架构")
    add_body(doc, "后端项目位于 backend/app 目录。main.py 负责创建 FastAPI 应用、注册中间件、路由和技能；routers 目录按业务域提供 REST API；services 目录承载业务逻辑；models 目录定义数据库模型；schemas 目录定义请求和响应结构；agents、skills、prompts、adapters、integrations 等目录承载 AI 与模型调用能力。")
    add_body(doc, "后端使用异步 SQLAlchemy 访问 SQLite 数据库，适合本地演示和轻量部署。项目已开启数据库 WAL 模式和启动备份机制，减少演示过程中的数据损坏风险。认证采用 token 机制，受保护接口通过当前用户依赖进行权限校验。")

    add_table(
        doc,
        ["后端模块", "职责说明"],
        [
            ["routers/auth.py", "登录、注册和当前用户认证。"],
            ["routers/needs.py", "需求发布、查询、删除、AI 描述生成、申请和匹配入口。"],
            ["routers/agents.py", "Agent 会话、消息、文件上传、任务状态、草稿生成和推荐能力。"],
            ["routers/messages.py", "站内消息发送、会话列表和按需求沟通。"],
            ["routers/profile.py", "用户资料、技能标签、扩展画像和标签提取。"],
            ["routers/settings.py", "模型 API Key 和服务设置管理。"],
        ],
    )

    add_heading(doc, "四、数据模型设计")
    add_body(doc, "平台核心数据包括用户、需求、匹配结果、申请、消息和 Agent 会话。用户模型保存用户名、技能标签、个人简介、学校和扩展画像；需求模型保存类型、标题、描述、发布者、单人或多人模式、状态和语义向量；匹配结果保存候选人、分数、理由和反馈；申请模型保存申请方、需求、申请消息和处理状态；消息模型保存发送方、接收方、需求上下文和内容；Agent 模型保存会话、消息、任务和上传文件。")
    add_body(doc, "数据设计强调业务状态一致性。单人需求完成配对后，会同步处理相关待处理申请，避免出现发布方已完成匹配而申请方仍显示待处理。删除需求时会处理关联匹配任务和申请数据，避免后端返回 500 或前端页面崩溃。")

    add_heading(doc, "五、AI 与 Agent 架构")
    add_body(doc, "AI 能力采用 PromptRegistry、SkillRegistry 和 Agent 组合模式。PromptRegistry 统一管理语义路由、意图判断、文件分析、需求生成、消息起草、匹配解释等提示词，避免在业务代码中散落长提示词。SkillRegistry 将语义路由、文件读取、标签提取、Embedding、向量匹配、匹配解释、内容审核、任务规划和上下文摘要封装为可复用能力。")
    add_body(doc, "Agent 层负责把用户输入转化为业务动作。IntentAnalyzerAgent 负责理解用户意图；PlannerAgent 负责规划任务；NeedCreatorAgent 负责生成需求草稿；SemanticSearchAgent 负责语义召回；RerankAgent 负责候选人精排和推荐理由生成；MatchWatcherAgent 负责匹配完成后的提醒和状态推进。")
    add_body(doc, "新增的 semantic_router skill 用于统一判断用户真实意图，输出 intent、next_action、semantic_frame、safety_level 和 rationale。该能力解决了“帮我找现有需求”被误判为“发布需求”的问题，也能在发布需求时识别缺失信息，要求继续追问技能方向和人数模式。")

    add_heading(doc, "六、模型与底座适配")
    add_body(doc, "平台将生成类任务和语义检索类任务分开处理。生成类任务主要通过 DeepSeek Chat 适配器完成，包括 Agent 对话、需求润色、文件分析、申请文案、私信文案和推荐理由生成。语义向量和重排任务通过本地 Qwen3-Embedding-0.6B 与 Qwen3-Reranker-0.6B 适配器完成，在模型不可用时会降级到 fallback 逻辑，保证主链路不中断。")
    add_body(doc, "模型调用通过 integrations/client.py 和 integrations/model_router.py 进行路由。系统设置页支持更新 API Key，新的 AI 请求能够读取最新配置。为了演示安全，前端和文档不应展示完整 API Key。")

    add_heading(doc, "七、匹配算法流程")
    add_body(doc, "匹配流程从需求发布开始。系统先抽取需求中的技能标签，再生成需求语义向量，并与用户画像、技能标签和历史信息进行匹配。第一阶段使用向量召回和标签相关性找到候选人；第二阶段使用 Qwen3 Reranker 对候选人进行精排；第三阶段使用 DeepSeek 生成可读的推荐理由和沟通建议。")
    add_body(doc, "匹配结果面向用户展示为候选人卡片和对比表，不只显示分数，还展示技能、简介、匹配原因和可联系入口。发布方可以选择候选人，系统根据单人或多人模式更新需求状态，并触发后续消息沟通或申请状态流转。")

    add_heading(doc, "八、接口与事件机制")
    add_body(doc, "后端通过 FastAPI 提供 REST API。前端 API 模块统一封装请求，开发环境下通过 Vite proxy 访问 /api，降低本地端口差异带来的配置成本。主要接口覆盖认证、个人资料、需求、申请、匹配、消息、Agent 会话、文件上传、系统设置和技能列表。")
    add_body(doc, "项目使用轻量 EventBus 作为 hooks 机制，在用户注册、需求发布、反馈接收、Agent 文件处理、Agent 创建需求和匹配完成等关键节点触发事件。事件机制使业务接口和后续扩展解耦，便于后续增加主动提醒、偏好更新和长期记忆。")

    add_heading(doc, "九、安全与稳定性设计")
    add_body(doc, "安全方面，平台对 Agent 会话、任务、匹配结果和需求操作进行了用户归属校验，防止用户查看或操作不属于自己的敏感数据。Agent 消息渲染使用文本插值而不是直接插入 HTML，降低 XSS 风险。系统设置页不应明文展示完整 API Key，演示材料中也要求遮挡敏感信息。")
    add_body(doc, "稳定性方面，平台处理了重复点击确认发布导致多条重复需求、删除有关联申请和匹配任务的需求时报 500、消息页路由复用不刷新、长会话消息顺序错误、草稿切换页面丢失等问题。关键流程均有后端契约测试和 Agent 烟测覆盖。")

    add_heading(doc, "十、部署与运行")
    add_body(doc, "本地演示默认后端地址为 http://127.0.0.1:8000，前端地址为 http://127.0.0.1:5173。后端启动命令为 conda run -n ark python -m uvicorn app.main:app --host 127.0.0.1 --port 8000，前端启动命令为 npm run dev -- --host 127.0.0.1 --port 5173。后端健康检查接口为 /api/health，正常返回内容为 {\"status\":\"ok\"}。")
    add_body(doc, "项目提交代码包时建议包含 frontend、backend、docs、tests、requirements.txt、package.json 等核心内容，并排除 app.db、备份文件、上传文件、node_modules、模型缓存和 API Key 等敏感或体积较大的运行产物。")

    add_heading(doc, "十一、测试验证")
    add_body(doc, "最近一次验证结果显示，后端 tests.test_project_contracts 共 48 项通过，tests.test_agent_smoke 共 19 项通过，python compileall app tests 通过，前端 npm run build 通过。构建中存在来自第三方依赖的 Rolldown 注释警告和 vendor chunk size warning，属于当前已知非阻塞项。")
    add_body(doc, "测试覆盖重点包括登录与账号切换、系统设置 API Key 更新、手动发布需求、Agent 发布需求、Agent 查找已有需求、文件上传生成多草稿并选择发布、匹配结果与私信、主动申请、接受申请、单人需求配对后申请终态、被选中用户可见事件、需求删除不报 500、消息页刷新和路由切换等场景。")

    add_heading(doc, "十二、扩展方向")
    add_body(doc, "架构上后续可以继续扩展为更完整的校园 AI 协作平台。短期可增加后端草稿表、WebSocket 实时通知、演示保护开关和更细粒度的状态审计。中期可加入长期记忆、偏好学习、合作反馈和匹配质量评估。长期可扩展到小程序端、院系级需求池、导师资源库和跨平台协作接口。")

    out = OUT_DIR / "技术架构文档_校园AI互助撮合平台_提交版.docx"
    doc.save(out)
    return out


def main():
    if "--validate" in sys.argv:
        files = [
            OUT_DIR / "项目说明书_校园AI互助撮合平台_提交版.docx",
            OUT_DIR / "技术架构文档_校园AI互助撮合平台_提交版.docx",
        ]
        for path in files:
            xml = zipfile.ZipFile(path).read("word/document.xml").decode("utf-8")
            print(path.name.encode("unicode_escape").decode("ascii"))
            print("has_simhei_title_font=", 'w:eastAsia="黑体"' in xml)
            print("has_simsun_body_font=", 'w:eastAsia="宋体"' in xml)
            print("has_times_new_roman=", "Times New Roman" in xml)
            print("has_fixed_22pt_line_spacing=", 'w:line="440"' in xml and 'w:lineRule="exact"' in xml)
            print("has_two_char_first_indent=", 'w:firstLine="480"' in xml)
        return

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generated = [create_project_manual(), create_architecture_doc()]
    for path in generated:
        print(path)


if __name__ == "__main__":
    main()
